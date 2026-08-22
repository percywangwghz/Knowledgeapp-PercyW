# -*- coding: utf-8 -*-
"""pdf-preprocess 与 ingest 接入的单测（纯脚本，直接 python test_pdf.py 跑）。

覆盖：语义切分/均匀切分、本地库链各档（真实研报 PDF 冒烟）、VLM 档消息组装与
串行退避（mock requests.post，不触网不烧 token）、单块全链失败降级、
docx 段落表格交错、organize_document 分段整理、入库不保留原文、任务占用提示。
"""
import io
import json
import os
import sys
import tempfile
import unittest.mock as mock

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# 测试环境固定端点为 Moonshot 预设（vision=True），不受本机 env 影响
os.environ.pop("KB_BASE_URL", None)
os.environ.pop("MOONSHOT_BASE_URL", None)

import config
import ingest

# 按文件路径加载 convert.py（目录名含连字符，同 ingest._load_convert）
convert = ingest._load_convert()

failures = []


def check(name, cond, detail=""):
    if cond:
        print(f"[OK]   {name}")
    else:
        failures.append(name)
        print(f"[FAIL] {name} {detail}")


def make_pdf(page_texts):
    """用 pymupdf 现场构造 PDF：每条文本一页，返回字节。"""
    import pymupdf
    doc = pymupdf.open()
    for text in page_texts:
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=12)
    buf = doc.tobytes()
    doc.close()
    return buf


REAL_PDF = os.path.join(os.path.dirname(os.path.abspath(__file__)), os.pardir,
                        "光模块", "20260612-浙商证券-通信行业专题报告：光模块技术路线梳理.pdf")

# ==================== 切分 ====================
# 语义切点：25 页，第 0/5/12 页（0-based）页首为章节标题
texts = [f"page {i} content" for i in range(25)]
texts[5] = "1. Industry Overview\npage 5 content"
texts[12] = "2. Market Size\npage 12 content"
breaks = convert._find_breaks(texts)
check("语义切点识别", {5, 12} <= breaks, f"breaks={breaks}")

chunks = convert._plan_chunks(25, breaks, 10)
check("语义切块边界", chunks == [(0, 5), (5, 12), (12, 22), (22, 25)], f"chunks={chunks}")
check("切块不超上限", all(e - s <= 10 for s, e in chunks))

chunks = convert._plan_chunks(25, set(), 10)
check("无切点均匀切", chunks == [(0, 10), (10, 20), (20, 25)], f"chunks={chunks}")
check("小 PDF 单块", convert._plan_chunks(3, set(), 10) == [(0, 3)])
# 目录点线行不算标题
toc_texts = ["Table of Contents\n1. Industry ...... 5", "page 1 content"]
check("目录点线不误判", convert._find_breaks(toc_texts) == set())

# ==================== 本地链（真实 PDF 冒烟，不传 key） ====================
if os.path.exists(REAL_PDF):
    with open(REAL_PDF, "rb") as f:
        real_data = f.read()
    prog = []
    out = convert.convert_pdf(os.path.basename(REAL_PDF), real_data,
                              api_key="", progress_cb=prog.append)
    engines = {b["engine"] for b in out["blocks"]}
    check("本地链走 pymupdf4llm", engines == {"pymupdf4llm"}, f"engines={engines}")
    check("本地链块全部 ok", all(b["ok"] for b in out["blocks"]))
    check("锚点注释注入", "<!-- p.1-" in out["md"] and "engine:pymupdf4llm" in out["md"])
    check("本地链产出标题结构", "#" in out["md"])
    check("本地链产出非平凡内容", len(out["md"]) > 5000, f"len={len(out['md'])}")
    check("无 key 不走 VLM", "vlm" not in engines)
    check("无降级页", out["degraded_pages"] == [])
    check("进度回调逐块", any(p.get("status") == "start" for p in prog)
          and any(p.get("status") == "done" for p in prog))
    print(f"       （本地链实测：{len(out['blocks'])} 块，{len(out['md'])} 字符，"
          f"表格 {'有' if '|' in out['md'] else '无'}）")
else:
    print(f"[SKIP] 真实 PDF 不存在：{REAL_PDF}")

# pdfplumber 档：禁用 pymupdf4llm 后落到该档
pdf3 = make_pdf([f"page {i} pdfplumber tier text" for i in range(3)])
with mock.patch.object(convert, "_chunk_pymupdf4llm",
                       side_effect=convert._EngineUnavailable("test 禁用")):
    out = convert.convert_pdf("t.pdf", pdf3, api_key="")
check("pdfplumber 档生效", all(b["engine"] == "pdfplumber" and b["ok"]
                               for b in out["blocks"]))
check("pdfplumber 档内容", "pdfplumber tier text" in out["md"])

# pypdf 档：两个本地库都禁用
with mock.patch.object(convert, "_chunk_pymupdf4llm",
                       side_effect=convert._EngineUnavailable("test 禁用")), \
     mock.patch.object(convert, "_chunk_pdfplumber",
                       side_effect=convert._EngineUnavailable("test 禁用")):
    out = convert.convert_pdf("t.pdf", pdf3, api_key="")
check("pypdf 档生效", all(b["engine"] == "pypdf" and b["ok"] for b in out["blocks"]))
check("pypdf 档不算降级", out["degraded_pages"] == [])

# ==================== 大 PDF 切分 + 锚点（60+ 页合成） ====================
big_texts = [f"body text of page {i}" for i in range(65)]
for i in (0, 11, 23, 37, 52):
    big_texts[i] = f"{i // 10 + 1}. Chapter Title {i}\nbody text of page {i}"
big_pdf = make_pdf(big_texts)
out = convert.convert_pdf("big.pdf", big_pdf, api_key="")
check("大 PDF 多块", len(out["blocks"]) >= 6, f"blocks={len(out['blocks'])}")
check("大 PDF 锚点逐块", all(f"<!-- p.{b['pages']} ·" in out["md"] for b in out["blocks"]))
check("大 PDF 语义切块", all(int(b["pages"].split("-")[1]) - int(b["pages"].split("-")[0]) < 10
                             for b in out["blocks"]),
      f"blocks={[b['pages'] for b in out['blocks']]}")
check("大 PDF 内容无丢失", "body text of page 64" in out["md"])

# ==================== VLM 档（mock requests.post，不触网） ====================
FAKE_MD = "# 转换结果\n\n| A | B |\n| --- | --- |\n| 1 | 2 |"


def fake_resp(text=FAKE_MD, status=200):
    r = mock.Mock()
    r.status_code = status
    r.text = text
    r.json.return_value = {"choices": [{"message": {"role": "assistant", "content": text}}]}
    return r


pdf_small = make_pdf([f"vlm page {i}" for i in range(3)])
with mock.patch("requests.post", return_value=fake_resp()) as post, \
     mock.patch.object(convert.time, "sleep"):
    out = convert.convert_pdf("v.pdf", pdf_small, api_key="fake-key", model="kimi-k2.6")
check("VLM 档生效", all(b["engine"] == "vlm" for b in out["blocks"]),
      f"blocks={out['blocks']}")
payload = post.call_args[1]["json"]
content = payload["messages"][0]["content"]
check("VLM 消息组装", payload["model"] == "kimi-k2.6"
      and content[0]["type"] == "text"
      and all(c["type"] == "image_url"
              and c["image_url"]["url"].startswith("data:image/png;base64,")
              for c in content[1:]))
check("VLM 页图数量=块页数", len(content) - 1 == 3)
check("VLM 输出进 md", FAKE_MD.split("\n")[0] in out["md"]
      and "engine:vlm" in out["md"])
check("VLM 超时元组（连接15s/读取180s）", convert.VLM_TIMEOUT == (15, 180)
      and post.call_args[1].get("timeout") == (15, 180))

# VLM 档块页数独立于本地档：VLM 按 KB_VLM_CHUNK_PAGES（默认 5）切，本地档仍 10
pdf12 = make_pdf([f"vlm chunk page {i}" for i in range(12)])
with mock.patch("requests.post", return_value=fake_resp()), \
     mock.patch.object(convert, "VLM_CHUNK_PAGES", 5), \
     mock.patch.object(convert.time, "sleep"):
    out = convert.convert_pdf("v.pdf", pdf12, api_key="fake-key", model="kimi-k2.6")
check("VLM 档按 5 页切块", [b["pages"] for b in out["blocks"]] == ["1-5", "6-10", "11-12"],
      f"blocks={[b['pages'] for b in out['blocks']]}")
check("VLM 档多块全 vlm", all(b["engine"] == "vlm" for b in out["blocks"]))
with mock.patch.object(convert, "VLM_CHUNK_PAGES", 5):
    out = convert.convert_pdf("v.pdf", pdf12, api_key="")
check("本地档仍按 10 页切", [b["pages"] for b in out["blocks"]] == ["1-10", "11-12"],
      f"blocks={[b['pages'] for b in out['blocks']]}")

# env KB_VLM_CHUNK_PAGES 生效（import 时读 env；另载一份全新模块实例验证，不动 convert）
def _fresh_convert():
    import importlib.util
    path = os.path.join(config.APP_DIR, "skills", "pdf-preprocess", "convert.py")
    spec = importlib.util.spec_from_file_location("kb_pdf_convert_env", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


os.environ["KB_VLM_CHUNK_PAGES"] = "7"
check("KB_VLM_CHUNK_PAGES env 生效", _fresh_convert().VLM_CHUNK_PAGES == 7)
os.environ.pop("KB_VLM_CHUNK_PAGES", None)
check("env 清除后还原默认 5", _fresh_convert().VLM_CHUNK_PAGES == 5)
os.environ["KB_VLM_WORKERS"] = "5"
check("KB_VLM_WORKERS env 生效", _fresh_convert().VLM_WORKERS == 5)
os.environ.pop("KB_VLM_WORKERS", None)
check("env 清除后并发默认 3", _fresh_convert().VLM_WORKERS == 3)

# 串行退避：429 × 2 后 200，共 3 次请求
with mock.patch("requests.post",
                side_effect=[fake_resp("x", 429), fake_resp("x", 429), fake_resp()]) as post, \
     mock.patch.object(convert.time, "sleep") as sleep:
    out = convert.convert_pdf("v.pdf", pdf_small, api_key="fake-key", model="kimi-k2.6")
check("VLM 429 退避重试", post.call_count == 3 and out["blocks"][0]["engine"] == "vlm")
check("退避间隔指数增长", [c[0][0] for c in sleep.call_args_list] == [2, 4],
      f"sleeps={[c[0][0] for c in sleep.call_args_list]}")

# VLM 持续 500 → 落本地链
with mock.patch("requests.post", side_effect=[fake_resp("x", 500)] * 9), \
     mock.patch.object(convert.time, "sleep"):
    out = convert.convert_pdf("v.pdf", pdf_small, api_key="fake-key", model="kimi-k2.6")
check("VLM 失败落本地链", all(b["engine"] == "pymupdf4llm" for b in out["blocks"]))

# 文本层分流：有文本层的块不请求 VLM（文本层优先，视觉模型只兜底扫描/图片页）
def make_dense_pdf(page_texts):
    """多行文本页（逐行插入）：单行超长会超出页宽被 pymupdf4llm 当页外内容丢弃。"""
    import pymupdf
    doc = pymupdf.open()
    for text in page_texts:
        page = doc.new_page()
        for li, line in enumerate(text.split("\n")):
            page.insert_text((72, 72 + li * 14), line, fontsize=12)
    buf = doc.tobytes()
    doc.close()
    return buf


dense = "\n".join(["body text layer line"] * 30)  # 630 字符/页 ≫ 阈值
long_pdf = make_dense_pdf([dense] * 6)
with mock.patch("requests.post", return_value=fake_resp()) as post, \
     mock.patch.object(convert.time, "sleep"):
    out = convert.convert_pdf("t.pdf", long_pdf, api_key="fake-key", model="kimi-k2.6")
check("有文本层不走 VLM", post.call_count == 0
      and all(b["engine"] == "pymupdf4llm" for b in out["blocks"]),
      f"engines={[b['engine'] for b in out['blocks']]}")

# 混合页：文本层块本地解析、扫描块走 VLM；并发执行后仍按原顺序重组
mixed_pdf = make_dense_pdf([dense] * 5 + [""] * 5)
with mock.patch("requests.post", return_value=fake_resp()), \
     mock.patch.object(convert.time, "sleep"):
    out = convert.convert_pdf("m.pdf", mixed_pdf, api_key="fake-key", model="kimi-k2.6")
check("扫描块才走 VLM", [b["engine"] for b in out["blocks"]] == ["pymupdf4llm", "vlm"],
      f"engines={[b['engine'] for b in out['blocks']]}")
check("并发后块顺序不乱", out["md"].find("<!-- p.1-5") < out["md"].find("<!-- p.6-10"))

# 非视觉模型（deepseek 预设 vision=False）不走 VLM
# （按 base_url 匹配厂家：把端点指到 deepseek，模型名也落 deepseek）
os.environ["KB_BASE_URL"] = "https://api.deepseek.com/v1"
with mock.patch("requests.post") as post:
    out = convert.convert_pdf("v.pdf", pdf_small, api_key="fake-key", model="deepseek-chat")
os.environ.pop("KB_BASE_URL", None)
check("非视觉模型跳过 VLM", post.call_count == 0
      and all(b["engine"] == "pymupdf4llm" for b in out["blocks"]))

# 单块全链失败 → degraded + pypdf 文字层保底（内容不丢）
with mock.patch.object(convert, "_chunk_vlm", side_effect=RuntimeError("vlm 挂")), \
     mock.patch.object(convert, "_chunk_pymupdf4llm", side_effect=RuntimeError("p4l 挂")), \
     mock.patch.object(convert, "_chunk_pdfplumber", side_effect=RuntimeError("pp 挂")), \
     mock.patch.object(convert, "_chunk_pypdf", side_effect=RuntimeError("pypdf 挂")), \
     mock.patch.object(convert.time, "sleep"):
    out = convert.convert_pdf("v.pdf", pdf_small, api_key="fake-key", model="kimi-k2.6")
check("全链失败降级标记", out["degraded_pages"] == [1, 2, 3], f"degraded={out['degraded_pages']}")
check("降级块记录错误", all(b["error"] for b in out["blocks"]))
check("降级仍保底文字层", "vlm page 0" in out["md"]
      and all(b["ok"] for b in out["blocks"]))

# ==================== 重要图表截取 ====================
TINY_PNG = __import__("base64").b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==")


def make_image_pdf():
    """第 1/2 页同一张大图（面积 26%，应去重只剩 1 张）+ 每页一枚小 logo（0.3%，应被滤掉）。"""
    import pymupdf
    doc = pymupdf.open()
    for _ in range(2):
        page = doc.new_page()
        page.insert_text((72, 72), "report page with chart", fontsize=12)
        page.insert_image(pymupdf.Rect(72, 100, 500, 400), stream=TINY_PNG)
        page.insert_image(pymupdf.Rect(72, 720, 120, 750), stream=TINY_PNG)
    buf = doc.tobytes()
    doc.close()
    return buf


out = convert.convert_pdf("img.pdf", make_image_pdf(), api_key="")
imgs = out["images"]
check("大图被截取且跨页去重", len(imgs) == 1 and imgs[0]["name"] == "p1_1.png",
      f"names={[im['name'] for im in imgs]}")
check("截图为 PNG 字节", bool(imgs) and imgs[0]["data"].startswith(b"\x89PNG"))

# 矢量图表页：40 条路径撑起大区域 → 整区截图（位图提取拿不到矢量图）
def make_vector_pdf():
    import pymupdf
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "vector chart page", fontsize=12)
    for i in range(40):
        x = 80 + i * 11
        page.draw_line(pymupdf.Point(x, 200), pymupdf.Point(x, 600))
        page.draw_rect(pymupdf.Rect(x, 200 + (i % 7) * 40, x + 8, 600))
    buf = doc.tobytes()
    doc.close()
    return buf


out = convert.convert_pdf("vec.pdf", make_vector_pdf(), api_key="")
check("矢量图表区截图", len(out["images"]) == 1
      and out["images"][0]["name"].startswith("p1_"),
      f"names={[im['name'] for im in out['images']]}")

# ==================== ingest 接入 ====================
orig_kb = config.KNOWLEDGE_DIR
config.KNOWLEDGE_DIR = tempfile.mkdtemp()

# extract_text：PDF 产出带锚点的 Markdown
text = ingest.extract_text("t.pdf", pdf3)
check("extract_text PDF 出 Markdown", text.startswith("<!-- p.1-3")
      and "pdfplumber tier text" in text)

# docx：段落与表格按文档顺序交错
import docx
d = docx.Document()
d.add_paragraph("开头段落")
t = d.add_table(rows=2, cols=2)
t.rows[0].cells[0].text = "表头A"
t.rows[0].cells[1].text = "表头B"
t.rows[1].cells[0].text = "值1"
t.rows[1].cells[1].text = "值2"
d.add_paragraph("结尾段落")
buf = io.BytesIO()
d.save(buf)
text = ingest.extract_text("t.docx", buf.getvalue())
i_head, i_tbl, i_tail = text.find("开头段落"), text.find("表头A | 表头B"), text.find("结尾段落")
check("docx 段落表格交错有序", 0 <= i_head < i_tbl < i_tail,
      f"head={i_head} tbl={i_tbl} tail={i_tail}\n{text!r}")

# organize_document 分段整理（map-reduce）
orig_limit = ingest.TEXT_LIMIT
ingest.TEXT_LIMIT = 200
long_text = "\n\n".join(f"## 小节{i}\n" + "正文内容" * 30 for i in range(6))
calls = []


def fake_chat(messages, **kw):
    calls.append(messages)
    if "分段整理」拼接" in messages[0]["content"]:  # MERGE_PROMPT 的标志性措辞
        return "合并后终稿"
    return f"整理后{len(calls)}"


with mock.patch.object(ingest, "chat", side_effect=fake_chat):
    stages = []
    r = ingest.organize_document("long.md", long_text, "", stage_cb=lambda s: stages.append(s))
check("长文多次调用", len(calls) > 1, f"calls={len(calls)}")
chunk_calls = [c for c in calls if "分段整理」拼接" not in c[0]["content"]]
merge_calls = [c for c in calls if "分段整理」拼接" in c[0]["content"]]
check("每块不超块大小", all(len(c[1]["content"]) < 200 + 200 for c in chunk_calls))
check("分段拼接后触发合并工序", len(merge_calls) == 1 and r["content"] == "合并后终稿")
check("合并草稿含各段产出", "整理后1" in merge_calls[0][1]["content"])
check("合并工序回调", stages == ["merge"], f"stages={stages}")
check("续篇不重复标题提示", any("续篇" in c[0]["content"] for c in chunk_calls))
check("框架标题提取", ingest._framework_headings("# 一、公司与团队\n正文\n## 1.1 业务\n") ==
      ["# 一、公司与团队", "## 1.1 业务"])
ingest.TEXT_LIMIT = orig_limit

# _dedupe_sections：同标题只留最长、白名单过滤重排、白名单不可信时只去重
draft = ("# T\n\n> 摘要\n\n## 1. A\n短\n\n## 2. B\n内容B\n\n## 1. A\n更长的内容AAAA\n\n## 9. 框架外\nxxx")
out = ingest._dedupe_sections(draft)
check("同标题留最长", out.count("## 1. A") == 1 and "更长的内容AAAA" in out and "框架外" in out)
wl = ["## 1. A", "## 2. B"]
out2 = ingest._dedupe_sections(draft, wl)
check("白名单过滤框架外", "框架外" not in out2 and out2.index("## 1. A") < out2.index("## 2. B"))
draft2 = "## 2. B\n内容B\n\n## 1. A\n内容A\n\n## 2. B\n内容B"
out3 = ingest._dedupe_sections(draft2, wl)
check("白名单重排", out3.index("## 1. A") < out3.index("## 2. B"))
out4 = ingest._dedupe_sections(draft2, ["## 完全无关X", "## 完全无关Y"])
check("白名单不可信只去重", "内容A" in out4 and "内容B" in out4)
# 编号变体互相匹配（框架「模块 1：核心投资哲学」对齐产出「1. Core Investment…」类）
out5 = ingest._dedupe_sections("## 附录 A：被投公司整理\nx",
                               ["### 附录 A：被投公司整理（认知的证据库）"])
check("附录编号变体匹配", "被投公司整理" in out5)

# _cap_digest_sections：背景速览硬上限 3 条，其余章节不动
digest_doc = ("# T\n\n## 1. A\n正文一\n正文二\n\n## 背景速览\n- 条1\n- 条2\n- 条3\n- 条4\n- 条5\n\n## 2. B\n正文B")
capped = ingest._cap_digest_sections(digest_doc)
check("背景速览截断到 3 条", "- 条3" in capped and "- 条4" not in capped)
check("速览外章节不动", "正文一" in capped and "正文二" in capped and "正文B" in capped)

# organize prompt 附图表清单（占位符机制）
calls2 = []


def fake_chat2(messages, **kw):
    calls2.append(messages)
    return "# T\n正文"


with mock.patch.object(ingest, "chat", side_effect=fake_chat2):
    ingest.organize_document("bp.pdf", "原文", "02_deals", images=["p3_1.png"])
check("图表清单入 prompt", "[[图:p3_1.png]]" in calls2[0][0]["content"]
      and "第3页" in calls2[0][0]["content"], f"prompt={calls2[0][0]['content'][-300:]!r}")

# 表格区域不截图：重叠率计算
import pymupdf as _pm
_r = _pm.Rect(0, 0, 100, 100)
check("表格重叠率", abs(convert._overlap_frac(_r, _pm.Rect(0, 0, 100, 60)) - 0.6) < 1e-6)
check("空重叠为 0", convert._overlap_frac(_r, _pm.Rect(200, 200, 300, 300)) == 0.0)

# _split_chunks 边界
check("无标题按字数切", len(ingest._split_chunks("x" * 500, 200)) == 3)
check("短文单块", ingest._split_chunks("短", 200) == ["短"])

# 入库不保留原文（用户明确要求）：折叠段已移除
result = {"category_key": "", "title": "T", "filename": "折叠测试",
          "summary": "s", "content": "# T\n\n正文", "text": "<!-- p.1-3 --> 全文内容"}
p = ingest.save_document(result, "来源.pdf")
saved = open(p, encoding="utf-8").read()
check("文末不带原文折叠段", "原文全文" not in saved and "全文内容" not in saved)

# 入库不再带文末图集：未被正文 [[图:…]] 引用的图（模型判为无关）不搬运、不入库
orig_data2 = ingest.DATA_DIR
ingest.DATA_DIR = tempfile.mkdtemp()
stage = os.path.join(ingest.DATA_DIR, ingest.IMG_STAGE_DIR, ingest._slug("来源.pdf"))
os.makedirs(stage)
with open(os.path.join(stage, "p3_1.png"), "wb") as f:
    f.write(b"\x89PNG fake")
res_img = {"file": "来源.pdf", "category_key": "01_industry", "title": "T",
           "filename": "图集测试", "summary": "s", "content": "# T\n\n正文",
           "text": "全文", "images": ["p3_1.png"]}
p = ingest.save_document(res_img, "来源.pdf")
saved = open(p, encoding="utf-8").read()
img_dst = os.path.join(config.KNOWLEDGE_DIR, "assets", "img", "图集测试", "p3_1.png")
check("未引用图不进文档", "其他图表" not in saved and "p3_1.png" not in saved,
      f"saved={saved[-400:]!r}")
check("未引用图不搬入 assets 且暂存清空", not os.path.isfile(img_dst)
      and not os.path.exists(stage))
ingest.DATA_DIR = orig_data2

# 占位符内联：被正文 [[图:…]] 引用的图搬进 assets 并原地替换，未引用的直接丢弃
ingest.DATA_DIR = tempfile.mkdtemp()
stage = os.path.join(ingest.DATA_DIR, ingest.IMG_STAGE_DIR, ingest._slug("来源.pdf"))
os.makedirs(stage)
for nm in ("p3_1.png", "p7_2.png"):
    with open(os.path.join(stage, nm), "wb") as f:
        f.write(b"\x89PNG fake")
res_tok = {"file": "来源.pdf", "category_key": "01_industry", "title": "T",
           "filename": "占位测试", "summary": "s",
           "content": "# T\n\n## 市场\n分析[[图:p3_1.png]]\n\n## 团队\n正文",
           "text": "全文", "images": ["p3_1.png", "p7_2.png"]}
p = ingest.save_document(res_tok, "来源.pdf")
saved = open(p, encoding="utf-8").read()
check("占位符原地替换", "分析![p.3 图表](../assets/img/占位测试/p3_1.png)" in saved,
      f"saved={saved[:400:]!r}")
check("引用图搬入 assets", os.path.isfile(
    os.path.join(config.KNOWLEDGE_DIR, "assets", "img", "占位测试", "p3_1.png")))
check("未引用图不出现也不搬运", "p7_2.png" not in saved
      and not os.path.isfile(os.path.join(config.KNOWLEDGE_DIR, "assets", "img",
                                          "占位测试", "p7_2.png")))
ingest.DATA_DIR = orig_data2

# _analyze_one 留存全文 + PDF 解析说明
with mock.patch.object(ingest, "chat",
                       side_effect=['{"category_key": "", "title": "T", "filename": "f", "summary": "s"}',
                                    "# T\n正文"]):
    res = ingest._analyze_one("t.pdf", pdf3)
check("_analyze_one 全文留存", res["ok"] and res["text"] == ingest.extract_text("t.pdf", pdf3))
check("_analyze_one 解析说明", "PDF 解析" in res.get("note", ""), f"note={res.get('note')}")

# 空白（扫描件模拟）PDF：判空忽略锚点注释
blank_pdf = make_pdf([""] * 3)
res = ingest._analyze_one("blank.pdf", blank_pdf)
check("空白 PDF 明确报错", not res["ok"] and "提取不到文本" in res["error"])

# 后台任务进度：PDF 文件的 detail 带块进度（块 x/y · 引擎），块计数落盘供前端 ETA
orig_data = ingest.DATA_DIR
ingest.DATA_DIR = tempfile.mkdtemp()
job_snaps = []
orig_write = ingest._write_job


def _spy(state, name=ingest.JOB_FILE):
    job_snaps.append(json.loads(json.dumps(state, ensure_ascii=False)))  # 快照防原地变异
    return orig_write(state, name)


with mock.patch.object(ingest, "chat",
                       side_effect=['{"category_key": "", "title": "T", "filename": "f", "summary": "s"}',
                                    "# T\n正文"]), \
     mock.patch.object(ingest, "_write_job", side_effect=_spy):
    ingest._run_job([("t.pdf", pdf3)], "test-key")
with open(os.path.join(ingest.DATA_DIR, "ingest_job.json"), encoding="utf-8") as f:
    job = json.load(f)
check("后台任务完成", job["status"] == "done" and job["results"][0]["ok"])
check("后台任务保留解析说明", "PDF 解析" in job["results"][0].get("note", ""))
_details = [s["steps"][0].get("detail", "") for s in job_snaps if s.get("steps")]
check("块进度 detail 新格式", any(d.startswith("块 1/1 · pymupdf4llm") for d in _details),
      f"details={_details}")
check("块计数落盘（前端信息条 ETA 用）",
      any((s.get("blocks") or {}).get("total") for s in job_snaps))
check("完成后瞬态字段清除", "block_started" not in job["steps"][0]
      and "partial" not in job["steps"][0])

# 改判分类后台入库（_run_save_job 同步直调）：流式 partial 节流落盘、完成后清除
plan = [{"file": "t.pdf", "ok": True, "category_key": "", "chosen": "09_tech",
         "title": "T", "filename": "改判测试", "summary": "s",
         "content": "# T\n正文", "text": "全文" * 100, "error": ""}]
save_snaps = []


def _spy_save(state, name=ingest.JOB_FILE):
    save_snaps.append(json.loads(json.dumps(state, ensure_ascii=False)))
    return orig_write(state, name)


def chat_save_chunks(messages, **kw):
    """模拟流式重整：20 个 chunk 后返回终稿。"""
    cb = kw.get("on_chunk")
    if cb:
        for n in range(20):
            cb(f"重整部分{n}")
    return "# T2\n重整正文"


with mock.patch.object(ingest, "chat", side_effect=chat_save_chunks) as c, \
     mock.patch.object(ingest, "_write_job", side_effect=_spy_save):
    ingest._run_save_job(plan, "test-key")
with open(os.path.join(ingest.DATA_DIR, "ingest_save_job.json"), encoding="utf-8") as f:
    sjob = json.load(f)
check("入库任务完成", sjob["status"] == "done" and sjob["results"][0]["ok"])
check("改判触发重整", c.call_count == 1 and sjob["results"][0]["category_key"] == "09_tech")
save_partial = [s for s in save_snaps if any("partial" in st for st in s.get("steps", []))]
check("入库 partial 节流落盘", 1 <= len(save_partial) <= 3,
      f"partial_writes={len(save_partial)}（共 20 次 on_chunk）")
check("入库完成后 partial 清除", "partial" not in sjob["steps"][0])
saved = open(sjob["results"][0]["path"], encoding="utf-8").read()
check("改判入库不保留原文", "重整正文" in saved and "原文全文" not in saved)
check("入库结果不带全文/正文", "text" not in sjob["results"][0]
      and "content" not in sjob["results"][0])
# 未改判但有原文仍整理（分析阶段只分类，整理统一在入库按确认框架执行）
plan[0]["chosen"] = plan[0]["category_key"] = "01_industry"
plan[0]["filename"] = "未改判测试"
with mock.patch.object(ingest, "chat", return_value="# T3\n整理正文") as c:
    ingest._run_save_job(plan, "test-key")
check("未改判有原文仍整理", c.call_count == 1)
# 无原文兜底：直接用已有正文落库（兼容旧版分析结果/外部计划）
plan[0]["text"] = ""
with mock.patch.object(ingest, "chat") as c:
    ingest._run_save_job(plan, "test-key")
check("无原文不重整", c.call_count == 0)
ingest.DATA_DIR = orig_data
config.KNOWLEDGE_DIR = orig_kb

# 任务占用：_start_job/_start_save_job 运行中返回 False
class _AliveThread:
    def is_alive(self):
        return True


ingest._job_thread = _AliveThread()
check("任务占用返回 False", ingest._start_job([("a.txt", b"x")]) is False)
ingest._job_thread = None
ingest._save_thread = _AliveThread()
check("入库占用返回 False", ingest._start_save_job([]) is False)
ingest._save_thread = None

print()
if failures:
    print(f"{len(failures)} test(s) failed")
    sys.exit(1)
print("ALL PDF TESTS OK")
