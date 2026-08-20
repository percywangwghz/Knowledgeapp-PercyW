# -*- coding: utf-8 -*-
"""ingest 单测：文本提取 + 分类/整理两阶段 + 框架注入 + 归档落盘（不触网）"""
import io
import json
import os
import sys
import tempfile
import time
import unittest.mock as mock
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import config
import ingest

orig_kb = config.KNOWLEDGE_DIR
failures = []


def check(name, cond, detail=""):
    if cond:
        print(f"[OK]   {name}")
    else:
        failures.append(name)
        print(f"[FAIL] {name} {detail}")


CLASSIFY_OUT = '{"category_key": "02_deals", "title": "某光引擎公司解剖", "filename": "某光引擎_解剖", "summary": "一句话"}'
CONTENT_MD = "# 某光引擎公司解剖\n\n> 摘要\n\n## 一、公司基本信息\n正文"

# ---------- 文本提取 ----------
check("txt utf-8", ingest.extract_text("a.txt", "光模块 订单".encode("utf-8")) == "光模块 订单")
check("txt gbk 回退", ingest.extract_text("a.txt", "光模块".encode("gbk")) == "光模块")
try:
    ingest.extract_text("a.pptx", b"x")
    check("不支持类型抛错", False)
except ValueError as e:
    check("不支持类型抛错", "不支持的文件类型" in str(e))

buf = io.BytesIO()
import docx
d = docx.Document()
d.add_paragraph("第一段：AI 算力")
d.add_paragraph("第二段：资本开支")
d.save(buf)
text = ingest.extract_text("a.docx", buf.getvalue())
check("docx 提取", "第一段" in text and "第二段" in text)


# ---------- _parse_json_obj ----------
check("JSON 解析", ingest._parse_json_obj(f"```json\n{CLASSIFY_OUT}\n```")["title"] == "某光引擎公司解剖")
check("非 JSON → {}", ingest._parse_json_obj("无法解析") == {})


# ---------- _sanitize_filename ----------
check("非法字符清洗", ingest._sanitize_filename('a/b\\c:d*e?"f<>|g h', "fb") == "a_b_c_d_e_f_g_h")
check("空名回退", ingest._sanitize_filename("  ", "原名") == "原名")
check("长度截断", len(ingest._sanitize_filename("x" * 100, "fb")) == ingest.MAX_FILENAME)


# ---------- 临时知识库（含框架文档） ----------
config.KNOWLEDGE_DIR = tempfile.mkdtemp()
fw_dir = os.path.join(config.KNOWLEDGE_DIR, "03_frameworks")
os.makedirs(fw_dir)
with open(os.path.join(fw_dir, "项目解剖模板.md"), "w", encoding="utf-8") as f:
    f.write("# 项目解剖模板\n## 三、假设树MARKER\n")
with open(os.path.join(fw_dir, "机构投资心智模型提取方法论.md"), "w", encoding="utf-8") as f:
    f.write("# 机构投资心智模型提取方法论\n## Investment Mind Model MARKER\n")


# ---------- classify_document（mock LLM） ----------
with mock.patch.object(ingest, "chat", return_value=CLASSIFY_OUT) as c:
    r = ingest.classify_document("某光引擎BP.pdf", "原文" * 5000)
check("classify 字段", r["category_key"] == "02_deals" and r["title"] == "某光引擎公司解剖"
      and r["filename"] == "某光引擎_解剖" and r["summary"] == "一句话")
check("classify 只看开头", len(c.call_args[0][0][1]["content"]) < 3000 + 200
      and c.call_args[1].get("feature") == "ingest")

with mock.patch.object(ingest, "chat", return_value=CLASSIFY_OUT.replace('"02_deals"', '"99_nowhere"')):
    r = ingest.classify_document("x.pdf", "原文")
check("非法分类兜底为根目录", r["category_key"] == "")

with mock.patch.object(ingest, "chat", return_value="完全无法解析的输出"):
    r = ingest.classify_document("原名.pdf", "原文")
check("无标题回退文件名", r["title"] == "原名" and r["filename"] == "原名")


# ---------- organize_document 框架注入 ----------
captured = {}


def fake_chat(messages, **kw):
    captured["messages"] = messages
    captured["kw"] = kw
    return f"```markdown\n{CONTENT_MD}\n```"


with mock.patch.object(ingest, "chat", side_effect=fake_chat):
    r = ingest.organize_document("bp.pdf", "原文", "02_deals")
prompt = captured["messages"][0]["content"]
check("框架注入 prompt", "假设树MARKER" in prompt and "研究框架" in prompt)
check("整理去外层 fence", r["content"].startswith("# 某光引擎公司解剖"))
check("organize feature 记账", captured["kw"].get("feature") == "ingest")

with mock.patch.object(ingest, "chat", side_effect=fake_chat):
    ingest.organize_document("fund.pdf", "原文", "08_funds")
check("基金框架注入", "Investment Mind Model MARKER" in captured["messages"][0]["content"])

with mock.patch.object(ingest, "chat", side_effect=fake_chat):
    ingest.organize_document("x.pdf", "原文", "07_learnings")
prompt = captured["messages"][0]["content"]
check("无框架走通用整理", "== 研究框架 ==" not in prompt and "划分章节" in prompt)

check("缺失框架文件返回空", ingest._load_framework("01_industry") == ""
      and ingest._load_framework("") == "")


# ---------- _analyze_one / ingest_one ----------
with mock.patch.object(ingest, "chat", side_effect=[CLASSIFY_OUT, CONTENT_MD]) as c:
    res = ingest._analyze_one("某光引擎BP.txt", "原文内容".encode("utf-8"))
check("_analyze_one 两次调用", c.call_count == 2)
check("_analyze_one 结果", res["ok"] and res["category_key"] == "02_deals"
      and res["content"].startswith("# 某光引擎公司解剖") and res["text"] and "path" not in res)

with mock.patch.object(ingest, "chat", return_value=CLASSIFY_OUT) as c:
    res = ingest._analyze_one("某光引擎BP.txt", "原文内容".encode("utf-8"), organize=False)
check("_analyze_one 仅分类模式一次调用", c.call_count == 1)
check("_analyze_one 仅分类不整理", res["ok"] and res["category_key"] == "02_deals"
      and not res["content"] and res["text"] == "原文内容")

with mock.patch.object(ingest, "chat", side_effect=[CLASSIFY_OUT, CONTENT_MD]) as c:
    res = ingest.ingest_one("某光引擎BP.txt", "原文内容".encode("utf-8"))
check("ingest_one 全流程", res["ok"] and os.path.exists(res["path"])
      and os.path.dirname(res["path"]).endswith("02_deals"))
check("ingest_one 分类→整理两次调用", c.call_count == 2)
check("ingest_one 整理直接吃原文",
      "原文内容" in c.call_args_list[1][0][0][1]["content"])
check("ingest_one 不带正文/原文", "text" not in res and "content" not in res)


# ---------- save_document ----------
result = {"category_key": "01_industry", "title": "光通信行业", "filename": "光通信_行业",
          "summary": "s", "content": "# 光通信行业\n\n## 小节\n正文"}
p1 = ingest.save_document(result, "来源.pdf")
check("归档到分类目录", os.path.dirname(p1).endswith("01_industry") and p1.endswith(".md"))
saved = open(p1, encoding="utf-8").read()
check("来源行插入标题后", saved.split("\n")[1].startswith("> 来源文件：来源.pdf · 归档于"))
p2 = ingest.save_document(result, "来源.pdf")
check("重名加序号", p2.endswith("光通信_行业_2.md"))
p3 = ingest.save_document({**result, "category_key": ""}, "来源.pdf")
check("空分类入根目录", os.path.dirname(p3) == config.KNOWLEDGE_DIR)

res = ingest.ingest_one("empty.txt", b"   ")
check("空文本报错", not res["ok"] and "提取不到文本" in res["error"])

# ---------- 后台任务（_run_job 同步直调，落盘验证；只解析+分类，整理在入库阶段） ----------
orig_data = ingest.DATA_DIR
ingest.DATA_DIR = tempfile.mkdtemp()
with mock.patch.object(ingest, "chat", side_effect=[CLASSIFY_OUT, CLASSIFY_OUT]):
    ingest._run_job([("a.txt", "原文A".encode()), ("b.txt", "原文B".encode()),
                     ("c.txt", b"   ")], "test-key")
with open(os.path.join(ingest.DATA_DIR, "ingest_job.json"), encoding="utf-8") as f:
    job = json.load(f)
check("后台任务落盘 done", job["status"] == "done" and len(job["results"]) == 3)
check("后台步骤状态", [s["status"] for s in job["steps"]] == ["done", "done", "error"])
check("后台结果含分类与原文、未整理", job["results"][0]["category_key"] == "02_deals"
      and job["results"][0]["text"] == "原文A" and not job["results"][0]["content"])
check("后台完成步 frac 收尾", job["steps"][0]["frac"] == 1.0)
ingest.DATA_DIR = orig_data

# ---------- on_chunk 透传（classify / organize） ----------
with mock.patch.object(ingest, "chat", return_value=CLASSIFY_OUT) as c:
    ingest.classify_document("x.pdf", "原文", on_chunk=lambda t: None)
check("classify 透传 on_chunk", c.call_args[1].get("on_chunk") is not None)

_sentinel = lambda t: None
with mock.patch.object(ingest, "chat", side_effect=fake_chat):
    ingest.organize_document("bp.pdf", "原文", "02_deals", on_chunk=_sentinel)
check("organize 透传 on_chunk", captured["kw"].get("on_chunk") is _sentinel)

# ---------- 打字机 partial：节流落盘 + 完成后清除 ----------
ingest.DATA_DIR = tempfile.mkdtemp()
snaps = []
orig_write = ingest._write_job


def _spy_write(state, name=ingest.JOB_FILE):
    snaps.append(json.loads(json.dumps(state, ensure_ascii=False)))  # 快照防原地变异
    return orig_write(state, name)


_seq = iter([CLASSIFY_OUT, CONTENT_MD])


def chat_with_chunks(messages, **kw):
    """模拟流式：每次 chat 触发 30 个 chunk 后返回。"""
    cb = kw.get("on_chunk")
    if cb:
        for n in range(30):
            cb(f"部分文本{n}")
    return next(_seq)


with mock.patch.object(ingest, "chat", side_effect=chat_with_chunks), \
     mock.patch.object(ingest, "_write_job", side_effect=_spy_write):
    ingest._run_job([("a.txt", "原文A".encode("utf-8"))], "test-key")
partial_snaps = [s for s in snaps if any("partial" in st for st in s.get("steps", []))]
check("partial 节流落盘出现", 1 <= len(partial_snaps) <= 5,
      f"partial_writes={len(partial_snaps)}（共 60 次 on_chunk）")
check("partial 内容为累计文本",
      bool(partial_snaps) and partial_snaps[0]["steps"][0]["partial"].startswith("部分文本"))
with open(os.path.join(ingest.DATA_DIR, "ingest_job.json"), encoding="utf-8") as f:
    job = json.load(f)
check("完成后 partial 清除", "partial" not in job["steps"][0] and job["status"] == "done")

# ---------- ETA 计算（mock 时间） ----------
start = int(time.time()) - 120
iso = datetime.fromtimestamp(start).isoformat()
sec = ingest._eta_seconds(iso, 2, 8, now=start + 120)
check("ETA 秒数换算", sec == 360.0, f"sec={sec}")  # elapsed120 / done2 * rem6
check("ETA 格式化分钟", ingest._fmt_eta(sec) == "约 6 分钟")
check("ETA 不足一分钟", ingest._fmt_eta(30) == "约 1 分钟内")
check("ETA 样本不足不估", ingest._eta_seconds(iso, 0, 8, now=start + 120) is None)
check("ETA 完成不估", ingest._eta_seconds(iso, 8, 8, now=start + 120) is None)
check("ETA 时间戳坏不估", ingest._eta_seconds("不是时间", 2, 8, now=start + 120) is None)

# ---------- 入库任务：文件粒度 ETA ----------
snaps.clear()
plan2 = [{"file": f"f{n}.txt", "ok": True, "category_key": "01_industry",
          "chosen": "01_industry", "title": f"T{n}", "filename": f"ETA测试{n}",
          "summary": "s", "content": f"# T{n}\n正文", "text": "", "error": ""}
         for n in (1, 2)]
with mock.patch.object(ingest, "_write_job", side_effect=_spy_write):
    ingest._run_save_job(plan2, "test-key")
first_run = next(s for s in snaps if s["steps"][0].get("status") == "running")
check("入库首文件无 ETA（样本不足）", not first_run["steps"][0].get("detail"),
      f"detail={first_run['steps'][0].get('detail')!r}")
d2 = [s["steps"][1].get("detail", "") for s in snaps
      if len(s.get("steps", [])) > 1 and s["steps"][1].get("status") == "running"]
check("入库按文件粒度 ETA", any("预计剩余" in d for d in d2), f"details={d2}")
ingest.DATA_DIR = orig_data

# ---------- 入库任务：凡有原文必按确认分类整理（唯一整理入口） ----------
snaps.clear()
ingest.DATA_DIR = tempfile.mkdtemp()
plan3 = [{"file": "f1.txt", "ok": True, "category_key": "02_deals",
          "chosen": "02_deals", "title": "T1", "filename": "入库整理1",
          "summary": "s", "content": "", "text": "原文一" * 100, "error": ""},
         {"file": "f2.txt", "ok": True, "category_key": "",
          "chosen": "02_deals", "title": "T2", "filename": "入库整理2",
          "summary": "s", "content": "", "text": "原文二" * 100, "error": ""}]
with mock.patch.object(ingest, "chat", return_value=CONTENT_MD) as c3, \
     mock.patch.object(ingest, "_write_job", side_effect=_spy_write):
    ingest._run_save_job(plan3, "test-key")
check("入库每文件一次整理调用", c3.call_count == 2)
check("整理用确认框架", "假设树MARKER" in c3.call_args_list[0][0][0][0]["content"])
with open(os.path.join(ingest.DATA_DIR, "ingest_save_job.json"), encoding="utf-8") as f:
    sjob = json.load(f)
check("入库任务 done 且落库", sjob["status"] == "done"
      and all(r.get("path") for r in sjob["results"]))
rejudge = [s["steps"][1].get("detail", "") for s in snaps
           if len(s.get("steps", [])) > 1 and s["steps"][1].get("status") == "running"]
check("改判 detail 标注", any("分类改判" in d for d in rejudge), f"details={rejudge}")
ingest.DATA_DIR = orig_data

# ---------- 入库任务：分析失败的文件不写空壳 ----------
ingest.DATA_DIR = tempfile.mkdtemp()
plan_bad = [{"file": "bad.txt", "ok": False, "category_key": "", "chosen": "",
             "title": "", "filename": "坏文件", "summary": "", "content": "",
             "text": "", "error": "提取不到文本"}]
ingest._run_save_job(plan_bad, "test-key")
with open(os.path.join(ingest.DATA_DIR, "ingest_save_job.json"), encoding="utf-8") as f:
    sjob = json.load(f)
check("失败文件不落盘带 error", not sjob["results"][0].get("path")
      and sjob["results"][0]["error"] == "提取不到文本")
check("失败文件不生成空壳文档",
      not os.path.exists(os.path.join(config.KNOWLEDGE_DIR, "坏文件.md")))
ingest.DATA_DIR = orig_data

# ---------- 上传缓存（重启后「重新发起」的数据基础） ----------
ingest.DATA_DIR = tempfile.mkdtemp()
ingest._cache_uploads([("a.txt", "原文A".encode("utf-8")), ("b.pdf", b"%PDF-x")])
cached = ingest._load_cached_uploads()
check("上传缓存回读文件数", len(cached) == 2, f"got={len(cached)}")
check("上传缓存回读内容一致", dict(cached).get("b.pdf") == b"%PDF-x",
      f"got={dict(cached).get('b.pdf')!r}")
ingest._cache_uploads([("c.md", "新批次".encode("utf-8"))])  # 新批次应清掉上一批
names = [n for n, _ in ingest._load_cached_uploads()]
check("新批次覆盖旧缓存", names == ["c.md"], f"got={names}")
ingest._clear_upload_cache()
check("清除缓存后回读为空", ingest._load_cached_uploads() == [],
      f"got={ingest._load_cached_uploads()!r}")
ingest.DATA_DIR = orig_data

config.KNOWLEDGE_DIR = orig_kb

print()
if failures:
    print(f"{len(failures)} test(s) failed")
    sys.exit(1)
print("ALL INGEST TESTS OK")
